"""
Document Processing Module
Handles file extraction, text cleaning, and experience extraction
"""

import re
import pdfplumber
import docx
import streamlit as st


def extract_text(file):
    """Extract text from PDF, DOCX, or TXT files - handles multi-page documents"""
    try:
        if file.type == "application/pdf":
            text = ""
            file.seek(0)
            
            try:
                with pdfplumber.open(file) as pdf:
                    total_pages = len(pdf.pages)
                    st.info(f"Processing PDF: {total_pages} page(s) detected")
                    
                    for page_num, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"
                        else:
                            st.warning(f"Page {page_num} appears to be empty or contains only images")
                    
                    if not text.strip():
                        st.error("No text could be extracted from the PDF. The file might contain only images or be corrupted.")
                        return ""
            except Exception as pdf_error:
                st.error(f"**PDF Processing Error:** Unable to read the PDF file.")
                st.warning(f"Error details: {str(pdf_error)}")
                st.info("**Possible solutions:**\n- The PDF might be corrupted or password-protected\n- Try opening and re-saving the PDF\n- Convert it to DOCX or TXT format")
                return ""
                    
            return text.strip()

        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            file.seek(0)
            
            try:
                document = docx.Document(file)
                
                paragraphs = [para.text for para in document.paragraphs if para.text.strip()]
                
                table_text = []
                for table in document.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            table_text.append(" | ".join(row_text))
                
                all_text = "\n".join(paragraphs)
                if table_text:
                    all_text += "\n" + "\n".join(table_text)
                
                if not all_text.strip():
                    st.error("No text could be extracted from the DOCX file.")
                    st.info("The document might be empty or contain only images.")
                    return ""
                
                total_paragraphs = len(paragraphs) + len(table_text)
                st.info(f"Processing DOCX: {total_paragraphs} paragraph(s)/section(s) detected")
                
                return all_text.strip()
            except Exception as docx_error:
                st.error(f"**DOCX Processing Error:** Unable to read the Word document.")
                st.warning(f"Error details: {str(docx_error)}")
                st.info("**Possible solutions:**\n- The file might be corrupted\n- Try opening and re-saving the document\n- Save as .docx format (not .doc)\n- Convert to PDF or TXT format")
                return ""

        elif file.type == "text/plain":
            file.seek(0)
            
            try:
                text = file.read().decode("utf-8")
                
                if not text.strip():
                    st.error("The text file appears to be empty.")
                    return ""
                
                line_count = len([line for line in text.split('\n') if line.strip()])
                st.info(f"Processing TXT: {line_count} line(s) detected")
                
                return text.strip()
            except UnicodeDecodeError:
                try:
                    file.seek(0)
                    text = file.read().decode("latin-1")
                    st.warning("File encoding detected as Latin-1 instead of UTF-8")
                    return text.strip()
                except Exception as encoding_error:
                    st.error(f"**Encoding Error:** Unable to read the text file.")
                    st.warning(f"Error details: {str(encoding_error)}")
                    st.info("**Possible solutions:**\n- Save the file with UTF-8 encoding\n- Copy content to a new text file\n- Convert to PDF or DOCX format")
                    return ""
            except Exception as txt_error:
                st.error(f"**TXT Processing Error:** Unable to read the text file.")
                st.warning(f"Error details: {str(txt_error)}")
                st.info("The file might be corrupted or in an unsupported format.")
                return ""
        else:
            st.error(f"**Unsupported file type:** {file.type}")
            st.warning("Please upload only PDF, DOCX, or TXT files.")
            return ""
        
    except Exception as e:
        st.error(f"**Unexpected Error** while processing {file.name}")
        st.warning(f"Error details: {str(e)}")
        st.info("**Try:**\n- Re-uploading the file\n- Using a different file format\n- Checking if the file is corrupted")
        return ""


def clean_text(text):
    """Clean and normalize text for NLP processing - optimized for longer documents"""
    if not text:
        return ""
    
    text = text.lower()
    
    # Preserve experience ranges before cleaning
    text = re.sub(r'(\d+)\s*(?:-|–|to)\s*(\d+)', r'\1_\2', text)
    text = re.sub(r'(\d+)\s*\+', r'\1_plus', text)
    
    text = re.sub(r'\n+', ' ', text)
    text = re.sub(r'\t+', ' ', text)
    
    text = re.sub(r'[^a-z0-9_\s]', ' ', text)
    
    text = re.sub(r'\s+', ' ', text).strip()
    
    return text


def extract_experience(text):
    """Extract years of experience from text - improved pattern matching"""
    if not text:
        return None
    text = text.lower()
    
    # Pattern 1: Experience range
    range_match = re.search(r'(\d+)\s*(?:-|–|to)\s*(\d+)\s*(?:years?|yrs?)', text)
    if range_match:
        return {"min_exp": int(range_match.group(1)), "max_exp": int(range_match.group(2))}
    
    # Pattern 2: Experience plus
    plus_match = re.search(r'(\d+)\s*\+\s*(?:years?|yrs?)', text)
    if plus_match:
        return {"min_exp": int(plus_match.group(1)), "max_exp": None}
    
    # Pattern 3: "with X years"
    with_years_match = re.search(r'(?:with|over|around|approximately)\s+(\d+)\s+(?:years?|yrs?)', text)
    if with_years_match:
        years = int(with_years_match.group(1))
        return {"min_exp": years, "max_exp": years}
    
    # Pattern 4: "X years of experience"
    years_of_match = re.search(r'(\d+)\s+(?:years?|yrs?)\s+(?:of\s+)?(?:experience|specializing|in|working)', text)
    if years_of_match:
        years = int(years_of_match.group(1))
        return {"min_exp": years, "max_exp": years}
    
    # Pattern 5: General fallback
    general_match = re.search(r'(\d+)\s+(?:years?|yrs?)', text)
    if general_match:
        years = int(general_match.group(1))
        return {"min_exp": years, "max_exp": years}
    
    return None
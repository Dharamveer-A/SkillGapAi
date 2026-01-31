"""
Report Generator Module
Generates reports in CSV, DOCX, and PDF formats
"""

import io
import pandas as pd
from docx import Document
from fpdf import FPDF
import math
import datetime

def generate_csv_report(skill_match_result, jd_skills):
    """Generate CSV report separating technical and soft skills"""
    tech_matched = [s for s in skill_match_result['matched'] if s in jd_skills["technical"]]
    tech_partial = [s for s in skill_match_result['partial'] if s in jd_skills["technical"]]
    tech_missing = [s for s in skill_match_result['missing'] if s in jd_skills["technical"]]
    
    soft_matched = [s for s in skill_match_result['matched'] if s in jd_skills["soft"]]
    soft_partial = [s for s in skill_match_result['partial'] if s in jd_skills["soft"]]
    soft_missing = [s for s in skill_match_result['missing'] if s in jd_skills["soft"]]
    
    max_tech_len = max(len(tech_matched), len(tech_partial), len(tech_missing), 1)
    tech_matched += [''] * (max_tech_len - len(tech_matched))
    tech_partial += [''] * (max_tech_len - len(tech_partial))
    tech_missing += [''] * (max_tech_len - len(tech_missing))
    
    max_soft_len = max(len(soft_matched), len(soft_partial), len(soft_missing), 1)
    soft_matched += [''] * (max_soft_len - len(soft_matched))
    soft_partial += [''] * (max_soft_len - len(soft_partial))
    soft_missing += [''] * (max_soft_len - len(soft_missing))
    
    tech_df = pd.DataFrame({
        "Technical - Matched": tech_matched,
        "Technical - Partially Matched": tech_partial,
        "Technical - Missing": tech_missing
    })
    
    soft_df = pd.DataFrame({
        "Soft Skills - Matched": soft_matched,
        "Soft Skills - Partially Matched": soft_partial,
        "Soft Skills - Missing": soft_missing
    })
    
    csv_output = tech_df.to_csv(index=False)
    csv_output += "\n"
    csv_output += soft_df.to_csv(index=False)
    
    return csv_output


def generate_word_report(skill_match_result, jd_skills, overall_match, matched_skills_count, match_counts):
    """Generate DOCX report"""
    doc = Document()
    title = doc.add_heading('Skill Gap Analysis Report', 0)
    title.alignment = 1
    doc.add_heading('Executive Summary', 1)
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = 'Table Grid' # Changed to standard grid for better compatibility
    summary_data = [
        ('Overall Match Rate', f'{overall_match}%'),
        ('Matched Skills', str(matched_skills_count)),
        ('Partially Matched Skills', str(match_counts["partial"])),
        ('Missing Skills', str(match_counts["missing"]))
    ]
    for i, (label, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value
    doc.add_paragraph()
    
    # Technical Skills Section
    doc.add_heading('Technical Skills Breakdown', 1)
    
    tech_matched = [s for s in skill_match_result['matched'] if s in jd_skills["technical"]]
    tech_partial = [s for s in skill_match_result['partial'] if s in jd_skills["technical"]]
    tech_missing = [s for s in skill_match_result['missing'] if s in jd_skills["technical"]]
    
    doc.add_heading('Matched Technical Skills', 2)
    if tech_matched:
        for skill in tech_matched:
            doc.add_paragraph(skill, style='List Bullet')
    else:
        doc.add_paragraph('No perfectly matched technical skills found.')
    
    doc.add_heading('Partially Matched Technical Skills', 2)
    if tech_partial:
        for skill in tech_partial:
            doc.add_paragraph(skill, style='List Bullet')
    else:
        doc.add_paragraph('No partially matched technical skills found.')
    
    doc.add_heading('Missing Technical Skills', 2)
    if tech_missing:
        for skill in tech_missing:
            doc.add_paragraph(skill, style='List Bullet')
    else:
        doc.add_paragraph('No missing technical skills!')
    
    # Soft Skills Section
    doc.add_heading('Soft Skills Breakdown', 1)
    
    soft_matched = [s for s in skill_match_result['matched'] if s in jd_skills["soft"]]
    soft_partial = [s for s in skill_match_result['partial'] if s in jd_skills["soft"]]
    soft_missing = [s for s in skill_match_result['missing'] if s in jd_skills["soft"]]
    
    doc.add_heading('Matched Soft Skills', 2)
    if soft_matched:
        for skill in soft_matched:
            doc.add_paragraph(skill, style='List Bullet')
    else:
        doc.add_paragraph('No perfectly matched soft skills found.')
    
    doc.add_heading('Partially Matched Soft Skills', 2)
    if soft_partial:
        for skill in soft_partial:
            doc.add_paragraph(skill, style='List Bullet')
    else:
        doc.add_paragraph('No partially matched soft skills found.')
    
    doc.add_heading('Missing Soft Skills', 2)
    if soft_missing:
        for skill in soft_missing:
            doc.add_paragraph(skill, style='List Bullet')
    else:
        doc.add_paragraph('No missing soft skills!')
    
    # Recommendations
    doc.add_heading('Recommendations', 1)
    doc.add_paragraph(
        f'Your current match rate is {overall_match}%. '
        f'To improve your candidacy, consider focusing on the missing skills listed above.'
    )
    
    if overall_match >= 70:
        doc.add_paragraph('You have a strong skill match for this position!', style='List Bullet')
    else:
        doc.add_paragraph('Focus on upskilling in the missing areas to increase your match rate.', style='List Bullet')
    
    footer = doc.add_paragraph(f'Generated on: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}')
    footer.alignment = 1
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

"""
Report Generator Module - Premium Edition
Generates a professional, two-column PDF report with visual skill tags and charts.
FIX: Solved 'module has no attribute now' error by using robust imports.
"""

from fpdf import FPDF
import math
import datetime  

class PDF(FPDF):
    def header(self):
        # 1. Top Purple Banner
        self.set_fill_color(71, 0, 71)  # #470047
        self.rect(0, 0, 210, 40, 'F')
        
        # 2. Title
        self.set_font('Arial', 'B', 24)
        self.set_text_color(255, 255, 255)
        self.set_xy(10, 10)
        self.cell(0, 15, 'Skill Gap Analysis Report', 0, 1, 'L')
        
        # 3. Subtitle / Confidentiality Marker
        self.set_font('Arial', 'I', 12)
        self.set_text_color(200, 200, 200) # Light grey
        self.cell(0, 10, 'Confidential Candidate Assessment', 0, 1, 'L')

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.set_text_color(128, 128, 128)
        
        # Safe timestamp generation
        date_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
        self.cell(0, 10, f'Generated by SkillGapAI on {date_str} | Page {self.page_no()}', 0, 0, 'C')

    def draw_section_title(self, title, x, y):
        """Draws a styled section header"""
        self.set_xy(x, y)
        self.set_font('Arial', 'B', 14)
        self.set_text_color(71, 0, 71)
        self.cell(0, 10, title, 0, 1, 'L')
        # Underline
        self.set_draw_color(200, 200, 200)
        self.line(x, y+9, x+80, y+9)

    def draw_skill_tag(self, text, x, y, type='match'):
        """
        Draws a colored rectangle with text inside (Visual 'Chip').
        """
        self.set_font('Arial', 'B', 9)
        w = self.get_string_width(text) + 6
        h = 7
        
        # Colors
        if type == 'match':
            self.set_fill_color(230, 250, 230) # Light Green BG
            self.set_text_color(0, 100, 0)     # Dark Green Text
            self.set_draw_color(46, 204, 113)  # Green Border
        elif type == 'partial':
            self.set_fill_color(255, 250, 230) # Light Orange BG
            self.set_text_color(150, 100, 0)   # Dark Orange Text
            self.set_draw_color(241, 196, 15)  # Orange Border
        else: # missing
            self.set_fill_color(250, 230, 230) # Light Red BG
            self.set_text_color(150, 0, 0)     # Dark Red Text
            self.set_draw_color(231, 76, 60)   # Red Border

        # Draw Tag
        self.rect(x, y, w, h, 'DF')
        self.set_xy(x, y)
        self.cell(w, h, text, 0, 0, 'C')
        
        return w + 2 # Return width + padding

    def draw_donut_chart(self, matched, partial, missing, x, y, radius, display_score):
        """Draws the vector Donut Chart"""
        total = matched + partial + missing
        if total == 0: return

        # Data: Green, Orange, Red
        data = [
            {'value': matched, 'color': (46, 204, 113)},   
            {'value': partial, 'color': (241, 196, 15)},   
            {'value': missing, 'color': (231, 76, 60)}     
        ]
        
        start_angle = 0
        
        # Draw Slices
        for item in data:
            if item['value'] > 0:
                angle = (item['value'] / total) * 360
                self.set_fill_color(*item['color'])
                self.set_line_width(0.1)
                
                # Draw smooth arc
                steps = int(angle * 2) 
                if steps < 2: steps = 2
                for i in range(steps):
                    a1 = math.radians(start_angle + (i * angle / steps))
                    a2 = math.radians(start_angle + ((i + 1) * angle / steps))
                    x1 = x + radius * math.cos(a1); y1 = y + radius * math.sin(a1)
                    x2 = x + radius * math.cos(a2); y2 = y + radius * math.sin(a2)
                    self.polygon([(x, y), (x1, y1), (x2, y2)], 'F')
                start_angle += angle

        # White Center
        inner_r = radius * 0.7
        self.draw_circle_custom(x, y, inner_r, (255, 255, 255))
        
        # Score Text
        self.set_font('Arial', 'B', 14)
        self.set_text_color(0, 0, 0)
        text = f"{display_score}%"
        text_w = self.get_string_width(text)
        self.set_xy(x - (text_w / 2), y - 2)
        self.cell(text_w, 4, text, 0, 0, 'C')

    def draw_circle_custom(self, x, y, r, color):
        self.set_fill_color(*color)
        self.set_line_width(0.1)
        steps = 60
        angle_step = 360 / steps
        for i in range(steps):
            a1 = math.radians(i * angle_step)
            a2 = math.radians((i + 1) * angle_step)
            x1 = x + r * math.cos(a1); y1 = y + r * math.sin(a1)
            x2 = x + r * math.cos(a2); y2 = y + r * math.sin(a2)
            self.polygon([(x, y), (x1, y1), (x2, y2)], 'F')

    def polygon(self, points, style='D'):
        op = {'D': 'S', 'F': 'f', 'FD': 'B', 'DF': 'B'}[style]
        s = '%.2f %.2f m ' % (points[0][0] * self.k, (self.h - points[0][1]) * self.k)
        for i in range(1, len(points)):
            s += '%.2f %.2f l ' % (points[i][0] * self.k, (self.h - points[i][1]) * self.k)
        s += ' h %s' % op
        self._out(s)

def generate_pdf_report(match_details, jd_skills, match_score, matched_count, metrics):
    pdf = PDF()
    pdf.add_page()

    # --- 1. Executive Summary & Chart ---
    pdf.draw_section_title("Executive Summary", 10, 50)
    pdf.set_font('Arial', '', 11)
    pdf.set_text_color(50, 50, 50)
    
    summary_text = (
        f"This automated analysis compares the provided resume against the job description.\n\n"
        f"OVERALL SCORE: {match_score}%\n"
        f"You have matched {matched_count} out of {metrics.get('total_jd', 0)} required skills.\n"
        f"Please review the detailed breakdown below to identify critical gaps."
    )
    pdf.set_xy(10, 65)
    pdf.multi_cell(110, 6, summary_text)
    
    # Draw Chart
    pdf.draw_donut_chart(metrics.get('matched', 0), metrics.get('partial', 0), metrics.get('missing', 0), 160, 75, 18, match_score)
    
    # --- 2. AI Verdict ---
    pdf.set_xy(10, 110)
    pdf.set_fill_color(240, 240, 245)
    pdf.rect(10, 110, 190, 25, 'F')
    
    pdf.set_xy(15, 115)
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 5, "AI Recommendation:", 0, 1)
    
    pdf.set_font('Arial', '', 11)
    pdf.set_x(15)
    if match_score >= 80:
        pdf.set_text_color(0, 100, 0)
        pdf.cell(0, 8, "Highly Recommended. This candidate shows a strong alignment with technical requirements.", 0, 1)
    elif match_score >= 50:
        pdf.set_text_color(150, 100, 0)
        pdf.cell(0, 8, "Potential Match. Candidate has core skills but lacks specific keywords found in the JD.", 0, 1)
    else:
        pdf.set_text_color(150, 0, 0)
        pdf.cell(0, 8, "Low Alignment. Significant skill gaps detected relative to the job requirements.", 0, 1)

    # --- 3. Two-Column Skill Analysis ---
    
    # Column 1: MATCHED (Left)
    y_start_cols = 150
    pdf.draw_section_title("Matched Skills", 10, y_start_cols) # Removed EMOJI
    
    curr_x = 10
    curr_y = y_start_cols + 15
    pdf.set_y(curr_y)
    
    if match_details['matched']:
        for skill in match_details['matched']:
            # Check if we need a new line
            if curr_x > 90: 
                curr_x = 10
                curr_y += 10
            
            # Draw Tag
            width = pdf.draw_skill_tag(skill, curr_x, curr_y, 'match')
            curr_x += width
    else:
        pdf.set_font('Arial', 'I', 10)
        pdf.text(10, curr_y + 5, "No exact matches found.")

    # Column 2: MISSING (Right)
    # We reset Y to the same starting height but move X to 110
    pdf.draw_section_title("Missing / Gaps", 110, y_start_cols) # Removed EMOJI
    
    curr_x = 110
    curr_y = y_start_cols + 15
    
    if match_details['missing']:
        for skill in match_details['missing']:
            # Check boundaries for right column
            if curr_x > 190: 
                curr_x = 110
                curr_y += 10
            
            # Draw Tag
            width = pdf.draw_skill_tag(skill, curr_x, curr_y, 'missing')
            curr_x += width
    else:
        pdf.set_xy(110, curr_y)
        pdf.set_font('Arial', 'I', 10)
        pdf.cell(0, 5, "No missing skills detected.")
        
    # --- 4. Partial Matches (Bottom) ---
    if match_details.get('partial'):
        final_y = max(curr_y, pdf.get_y()) + 25 # Move below the lowest column
        if final_y > 250: # Check if we need a new page
             pdf.add_page()
             final_y = 20
             
        pdf.draw_section_title("Partial / Related Matches", 10, final_y) # Removed EMOJI
        curr_x = 10
        curr_y = final_y + 15
        
        for skill in match_details['partial']:
            if curr_x > 180:
                curr_x = 10
                curr_y += 10
            width = pdf.draw_skill_tag(skill, curr_x, curr_y, 'partial')
            curr_x += width

    # Use 'replace' to safely handle any other rogue characters, though they should be gone
    return pdf.output(dest='S').encode('latin-1', 'replace')